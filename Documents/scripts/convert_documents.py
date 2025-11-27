#!/usr/bin/env python3
"""
Convert all binary documents to LLM-readable markdown format.

This script processes Excel, Word, and PDF files and creates markdown versions
that LLMs can easily consume while preserving the original files for download.
"""

import os
import json
from pathlib import Path
import pandas as pd
from docx import Document
import pypdf
from typing import Dict, Any

def excel_to_markdown(xlsx_path: Path) -> str:
    """Convert Excel file to markdown with tables for each sheet."""
    markdown = f"# {xlsx_path.stem.replace('_', ' ')}\n\n"
    markdown += f"*Source: {xlsx_path.name}*\n\n"
    
    try:
        xl = pd.ExcelFile(xlsx_path)
        
        for sheet_name in xl.sheet_names:
            df = pd.read_excel(xlsx_path, sheet_name=sheet_name)
            
            # Skip empty sheets
            if df.empty:
                continue
            
            markdown += f"## {sheet_name}\n\n"
            
            # Convert to markdown table
            markdown += df.to_markdown(index=False)
            markdown += "\n\n"
            
            # Add summary stats if numeric data
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                markdown += "### Summary Statistics\n\n"
                summary = df[numeric_cols].describe()
                markdown += summary.to_markdown()
                markdown += "\n\n"
        
        return markdown
    
    except Exception as e:
        return f"# {xlsx_path.stem}\n\n**Error converting Excel file:** {str(e)}\n"


def docx_to_markdown(docx_path: Path) -> str:
    """Convert Word document to markdown."""
    markdown = f"# {docx_path.stem.replace('_', ' ')}\n\n"
    markdown += f"*Source: {docx_path.name}*\n\n"
    
    try:
        doc = Document(docx_path)
        
        for para in doc.paragraphs:
            # Skip empty paragraphs
            if not para.text.strip():
                continue
            
            # Handle headings
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '')
                try:
                    level_num = int(level)
                    markdown += f"{'#' * (level_num + 1)} {para.text}\n\n"
                except ValueError:
                    markdown += f"## {para.text}\n\n"
            
            # Handle list items
            elif para.style.name.startswith('List'):
                markdown += f"- {para.text}\n"
            
            # Regular paragraphs
            else:
                markdown += f"{para.text}\n\n"
        
        # Handle tables
        if doc.tables:
            markdown += "\n## Tables\n\n"
            for i, table in enumerate(doc.tables):
                markdown += f"### Table {i+1}\n\n"
                
                # Convert table to list of lists
                data = []
                for row in table.rows:
                    data.append([cell.text for cell in row.cells])
                
                if data:
                    # Create markdown table
                    df = pd.DataFrame(data[1:], columns=data[0])
                    markdown += df.to_markdown(index=False)
                    markdown += "\n\n"
        
        return markdown
    
    except Exception as e:
        return f"# {docx_path.stem}\n\n**Error converting Word file:** {str(e)}\n"


def pdf_to_markdown(pdf_path: Path) -> str:
    """Convert PDF to markdown (basic text extraction)."""
    markdown = f"# {pdf_path.stem.replace('_', ' ')}\n\n"
    markdown += f"*Source: {pdf_path.name}*\n\n"
    
    try:
        reader = pypdf.PdfReader(pdf_path)
        
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            
            if text.strip():
                markdown += f"## Page {page_num}\n\n"
                markdown += text
                markdown += "\n\n---\n\n"
        
        return markdown
    
    except Exception as e:
        return f"# {pdf_path.stem}\n\n**Error converting PDF file:** {str(e)}\n"


def yaml_to_markdown(yaml_path: Path) -> str:
    """Copy YAML as-is (already readable) but add header."""
    markdown = f"# {yaml_path.stem.replace('_', ' ')}\n\n"
    markdown += f"*Source: {yaml_path.name}*\n\n"
    markdown += "```yaml\n"
    markdown += yaml_path.read_text(encoding='utf-8')
    markdown += "\n```\n"
    return markdown


def convert_all_documents():
    """Main conversion function."""
    base_dir = Path(__file__).parent.parent
    docs_dir = base_dir / "assets" / "documents"
    resources_dir = base_dir / "resources"
    markdown_dir = base_dir / "assets" / "markdown"
    resources_markdown_dir = base_dir / "resources" / "markdown"
    summaries_dir = base_dir / "assets" / "summaries"
    resources_summaries_dir = base_dir / "resources" / "summaries"
    
    # Create directories
    markdown_dir.mkdir(exist_ok=True, parents=True)
    resources_markdown_dir.mkdir(exist_ok=True, parents=True)
    summaries_dir.mkdir(exist_ok=True, parents=True)
    resources_summaries_dir.mkdir(exist_ok=True, parents=True)
    
    # Load document indexes
    converted_count = 0
    error_count = 0
    
    # Process assets/documents
    index_path = docs_dir / "document__index.json"
    if index_path.exists():
        print("=" * 60)
        print("Converting assets/documents...")
        print("=" * 60)
        with open(index_path, 'r', encoding='utf-8') as f:
            index = json.load(f)
        c, e = process_index(index, base_dir, markdown_dir, summaries_dir, index_path)
        converted_count += c
        error_count += e
    
    # Process resources
    resources_index_path = resources_dir / "document_index.json"
    if resources_index_path.exists():
        print("\n" + "=" * 60)
        print("Converting resources...")
        print("=" * 60)
        with open(resources_index_path, 'r', encoding='utf-8') as f:
            resources_index = json.load(f)
        c, e = process_index(resources_index, base_dir, resources_markdown_dir, resources_summaries_dir, resources_index_path)
        converted_count += c
        error_count += e
    
    print(f"\n{'='*60}")
    print(f"✅ All conversions complete!")
    print(f"   Total converted: {converted_count} documents")
    print(f"   Total errors: {error_count}")
    print(f"{'='*60}\n")


def process_index(index: dict, base_dir: Path, markdown_dir: Path, summaries_dir: Path, index_path: Path) -> tuple:
    """Process a single document index."""
    converted_count = 0
    error_count = 0
    
    converted_count = 0
    error_count = 0
    
    for doc in index['documents']:
        # Handle both 'id' and 'title' as identifier
        doc_id = doc.get('id')
        if not doc_id:
            # Generate ID from title for resources
            doc_id = doc['title'].lower().replace(' ', '_').replace('(', '').replace(')', '').replace(',', '').replace('–', '').replace('-', '_')
            doc['id'] = doc_id
        
        file_path = base_dir / doc['file']
        
        if not file_path.exists():
            print(f"⚠️  File not found: {file_path}")
            error_count += 1
            continue
        
        print(f"Converting: {file_path.name}...")
        
        # Convert based on file type
        if file_path.suffix == '.xlsx':
            markdown = excel_to_markdown(file_path)
        elif file_path.suffix == '.docx':
            markdown = docx_to_markdown(file_path)
        elif file_path.suffix == '.pdf':
            markdown = pdf_to_markdown(file_path)
        elif file_path.suffix in ['.yaml', '.txt']:
            markdown = yaml_to_markdown(file_path)
        else:
            print(f"  ⏭️  Skipping unsupported format: {file_path.suffix}")
            continue
        
        # Save markdown version
        md_path = markdown_dir / f"{doc_id}.md"
        md_path.write_text(markdown, encoding='utf-8')
        
        # Create summary (first 500 chars)
        summary = markdown[:500] + "..." if len(markdown) > 500 else markdown
        summary_path = summaries_dir / f"{doc_id}_summary.md"
        summary_path.write_text(summary, encoding='utf-8')
        
        # Determine relative paths based on which directory we're in
        if 'resources' in str(markdown_dir):
            md_rel = f"resources/markdown/{doc_id}.md"
            sum_rel = f"resources/summaries/{doc_id}_summary.md"
        else:
            md_rel = f"assets/markdown/{doc_id}.md"
            sum_rel = f"assets/summaries/{doc_id}_summary.md"
        
        # Update document entry with formats
        doc['formats'] = {
            'original': doc['file'],
            'markdown': md_rel,
            'summary': sum_rel
        }
        doc['llm_readable'] = True
        doc['user_viewable'] = True
        
        converted_count += 1
        print(f"  ✅ Converted to {md_path.name}")
    
    # Save updated index
    with open(index_path, 'w', encoding='utf-8') as f:
        json.dump(index, f, indent=2, ensure_ascii=False)
    
    print(f"\n   Converted: {converted_count} documents")
    print(f"   Errors: {error_count}")
    print(f"   Updated: {index_path}")
    
    return converted_count, error_count


if __name__ == "__main__":
    convert_all_documents()

